import json
from app.repositories import curriculos_repository as repo


def _parse_dados(row: dict) -> dict:
    dados = row.get('dados')
    if isinstance(dados, str):
        row['dados'] = json.loads(dados)
    return row


def list_curriculos(user_id: int) -> list:
    return [dict(r) for r in repo.list_curriculos(user_id)]


def get_curriculo(curriculo_id: int, user_id: int) -> dict | None:
    row = repo.get_curriculo(curriculo_id, user_id)
    if not row:
        return None
    return _parse_dados(dict(row))


def save_curriculo(user_id: int, curriculo_id: int | None, titulo: str, dados: dict) -> dict:
    titulo = (titulo or 'Meu Currículo').strip()[:200]
    if not titulo:
        titulo = 'Meu Currículo'
    if curriculo_id:
        repo.update_curriculo(curriculo_id, user_id, titulo, dados)
        row = repo.get_curriculo(curriculo_id, user_id)
    else:
        row = repo.create_curriculo(user_id, titulo, dados)
    return _parse_dados(dict(row))


def delete_curriculo(curriculo_id: int, user_id: int) -> None:
    repo.delete_curriculo(curriculo_id, user_id)


def _s(text) -> str:
    return (text or '').strip()


_LATIN1_SUBS = {
    '•': '·',  # bullet → middle dot (Latin-1)
    '–': '-',        # en dash
    '—': '-',        # em dash
    '‘': "'",        # left single quote
    '’': "'",        # right single quote
    '“': '"',        # left double quote
    '”': '"',        # right double quote
    '…': '...',      # ellipsis
}


def _t(text) -> str:
    s = _s(text)
    for ch, repl in _LATIN1_SUBS.items():
        s = s.replace(ch, repl)
    return s.encode('latin-1', errors='replace').decode('latin-1')


BULLET = '· '


def gerar_pdf(dados: dict) -> bytes:
    from fpdf import FPDF, XPos, YPos

    L = 20
    W = 170
    R = L + W

    pdf = FPDF(unit='mm', format='A4')
    pdf.set_margins(L, 18, L)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    def secao(titulo):
        pdf.ln(2)
        pdf.set_font('Helvetica', 'B', 9.5)
        pdf.set_text_color(0, 0, 0)
        pdf.set_draw_color(136, 136, 136)
        pdf.set_line_width(0.3)
        pdf.cell(W, 7, _t(titulo), border=1, align='C',
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)

    # ── NOME ──
    nome = _t(dados.get('nome', 'NOME COMPLETO')).upper()
    pdf.set_font('Helvetica', 'B', 17)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(W, 10, nome, new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # ── CABEÇALHO: linkedin/github esq · email/tel dir ──
    linkedin = _t(dados.get('linkedin'))
    github = _t(dados.get('github'))
    email = _t(dados.get('email'))
    telefone = _t(dados.get('telefone'))

    y_h = pdf.get_y()
    LEFT_W = 100
    RIGHT_W = W - LEFT_W

    # coluna direita
    pdf.set_xy(L + LEFT_W, y_h)
    if email:
        pdf.set_font('Helvetica', 'B', 9)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(RIGHT_W, 4.5, 'Email:', align='R',
                 new_x=XPos.LEFT, new_y=YPos.NEXT)
        pdf.set_x(L + LEFT_W)
        pdf.set_font('Helvetica', '', 9)
        pdf.set_text_color(51, 51, 51)
        pdf.cell(RIGHT_W, 4.5, email, align='R',
                 new_x=XPos.LEFT, new_y=YPos.NEXT)
        pdf.set_x(L + LEFT_W)
    if telefone:
        pdf.set_font('Helvetica', 'B', 9)
        pdf.set_text_color(0, 0, 0)
        pdf.cell(RIGHT_W, 4.5, 'Celular:', align='R',
                 new_x=XPos.LEFT, new_y=YPos.NEXT)
        pdf.set_x(L + LEFT_W)
        pdf.set_font('Helvetica', '', 9)
        pdf.set_text_color(51, 51, 51)
        pdf.cell(RIGHT_W, 4.5, telefone, align='R',
                 new_x=XPos.LEFT, new_y=YPos.NEXT)
    y_right_end = pdf.get_y()

    # coluna esquerda
    pdf.set_xy(L, y_h)
    pdf.set_font('Helvetica', '', 9)
    pdf.set_text_color(51, 51, 51)
    if linkedin:
        pdf.cell(LEFT_W, 4.5, 'LinkedIn: ' + linkedin,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    if github:
        pdf.cell(LEFT_W, 4.5, 'GitHub: ' + github,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    y_left_end = pdf.get_y()

    pdf.set_y(max(y_right_end, y_left_end))
    pdf.ln(3)

    # linha divisória
    y_line = pdf.get_y()
    pdf.set_draw_color(85, 85, 85)
    pdf.set_line_width(0.5)
    pdf.line(L, y_line, R, y_line)
    pdf.ln(6)

    # ── FORMAÇÃO ──
    formacao = [f for f in (dados.get('formacao') or [])
                if _s(f.get('instituicao')) or _s(f.get('descricao'))]
    if formacao:
        secao('FORMAÇÃO')
        RC = 55
        LC = W - RC - 3
        for f in formacao:
            inst = _t(f.get('instituicao'))
            desc = _t(f.get('descricao'))
            periodo = _t(f.get('periodo'))
            local = _t(f.get('localidade'))

            y_f = pdf.get_y()

            if periodo or local:
                pdf.set_xy(L + LC + 3, y_f)
                if periodo:
                    pdf.set_font('Helvetica', 'B', 9)
                    pdf.set_text_color(0, 0, 0)
                    pdf.cell(RC, 5, periodo, align='R',
                             new_x=XPos.LEFT, new_y=YPos.NEXT)
                    pdf.set_x(L + LC + 3)
                if local:
                    pdf.set_font('Helvetica', '', 9)
                    pdf.set_text_color(85, 85, 85)
                    pdf.cell(RC, 5, local, align='R',
                             new_x=XPos.LEFT, new_y=YPos.NEXT)
            y_right_f = pdf.get_y()

            pdf.set_xy(L, y_f)
            if inst:
                pdf.set_font('Helvetica', 'B', 9.5)
                pdf.set_text_color(0, 0, 0)
                pdf.multi_cell(LC, 5, inst)
                pdf.set_x(L)
            if desc:
                pdf.set_font('Helvetica', '', 9.5)
                pdf.set_text_color(51, 51, 51)
                pdf.multi_cell(LC, 5, desc)
            y_left_f = pdf.get_y()

            pdf.set_y(max(y_left_f, y_right_f))
            pdf.ln(4)

    # ── COMPETÊNCIAS ──
    comps = [c for c in (dados.get('competencias') or [])
             if _s(c.get('label')) or _s(c.get('valor'))]
    if comps:
        secao('RESUMO DE COMPETÊNCIAS')
        for c in comps:
            label = _t(c.get('label'))
            valor = _t(c.get('valor'))
            linha = (BULLET + label + ': ' + valor) if label and valor else (
                BULLET + (label or valor))
            pdf.set_x(L + 2)
            pdf.set_font('Helvetica', '', 9.5)
            pdf.set_text_color(51, 51, 51)
            pdf.multi_cell(W - 2, 5, linha)
            pdf.set_x(L)
        pdf.ln(2)

    # ── EXPERIÊNCIA ──
    exps = [e for e in (dados.get('experiencias') or [])
            if _s(e.get('cargo_contrato')) or _s(e.get('empresa'))]
    if exps:
        secao('EXPERIÊNCIA PROFISSIONAL')
        RC_E = 50
        LC_E = W - RC_E - 3
        for e in exps:
            cargo = _t(e.get('cargo_contrato'))
            empresa = _t(e.get('empresa'))
            periodo = _t(e.get('periodo'))
            resps = _t(e.get('responsabilidades'))
            title_str = cargo + (' | ' + empresa if empresa else '') if cargo else empresa

            y_e = pdf.get_y()

            if periodo:
                pdf.set_xy(L + LC_E + 3, y_e)
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(0, 0, 0)
                pdf.cell(RC_E, 5, periodo, align='R',
                         new_x=XPos.LEFT, new_y=YPos.NEXT)
            y_right_e = pdf.get_y()

            pdf.set_xy(L, y_e)
            pdf.set_font('Helvetica', 'B', 9.5)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(LC_E, 5, title_str)
            y_left_e = pdf.get_y()

            pdf.set_y(max(y_left_e, y_right_e))
            pdf.ln(1)

            if resps:
                bullets = [b.strip().lstrip('·-* ')
                           for b in resps.split('\n') if b.strip()]
                for b in bullets:
                    pdf.set_x(L + 4)
                    pdf.set_font('Helvetica', '', 9.5)
                    pdf.set_text_color(51, 51, 51)
                    pdf.multi_cell(W - 4, 5, BULLET + _t(b))
                    pdf.set_x(L)
            pdf.ln(4)

    # ── CERTIFICADOS ──
    certs = [c for c in (dados.get('certificados') or []) if _s(c.get('nome'))]
    if certs:
        secao('CERTIFICADOS')
        RC_C = 40
        LC_C = W - RC_C - 3
        for c in certs:
            nome_c = _t(c.get('nome'))
            data_c = _t(c.get('data'))

            y_c = pdf.get_y()
            if data_c:
                pdf.set_xy(L + LC_C + 3, y_c)
                pdf.set_font('Helvetica', 'B', 9)
                pdf.set_text_color(0, 0, 0)
                pdf.cell(RC_C, 5, data_c, align='R',
                         new_x=XPos.LEFT, new_y=YPos.NEXT)
            y_right_c = pdf.get_y()

            pdf.set_xy(L, y_c)
            pdf.set_font('Helvetica', '', 9.5)
            pdf.set_text_color(51, 51, 51)
            pdf.multi_cell(LC_C, 5, BULLET + nome_c)
            y_left_c = pdf.get_y()

            pdf.set_y(max(y_left_c, y_right_c))
            pdf.ln(2)

    # ── OUTROS ──
    import base64 as _b64
    from io import BytesIO as _BytesIO

    IMG_W = 35
    outros = [o for o in (dados.get('outros') or [])
              if _s(o.get('texto')) or (o.get('imagem') or '').startswith('data:image/')]
    if outros:
        secao('OUTROS')
        for o in outros:
            texto = _t(o.get('texto'))
            imagem_b64 = o.get('imagem') or ''
            has_img = imagem_b64.startswith('data:image/')

            if texto and has_img:
                TEXT_W = W - IMG_W - 5
                y_o = pdf.get_y()
                try:
                    raw = imagem_b64.split(',', 1)[1]
                    pdf.image(_BytesIO(_b64.b64decode(raw)),
                              x=L + TEXT_W + 5, y=y_o, w=IMG_W)
                except Exception:
                    has_img = False
                pdf.set_xy(L, y_o)
                pdf.set_font('Helvetica', '', 9.5)
                pdf.set_text_color(51, 51, 51)
                pdf.multi_cell(TEXT_W if has_img else W, 5, texto)
                pdf.set_y(max(pdf.get_y(), y_o + IMG_W))
                pdf.ln(4)

            elif texto:
                pdf.set_x(L)
                pdf.set_font('Helvetica', '', 9.5)
                pdf.set_text_color(51, 51, 51)
                pdf.multi_cell(W, 5, texto)
                pdf.ln(4)

            elif has_img:
                try:
                    raw = imagem_b64.split(',', 1)[1]
                    pdf.image(_BytesIO(_b64.b64decode(raw)),
                              x=L + (W - IMG_W) / 2, y=pdf.get_y(), w=IMG_W)
                    pdf.ln(IMG_W + 4)
                except Exception:
                    pass

    return bytes(pdf.output())


def gerar_docx(dados: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO

    doc = Document()

    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = Mm(20)
    sec.right_margin = Mm(20)
    sec.top_margin = Mm(18)
    sec.bottom_margin = Mm(18)

    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)

    def _no_borders(table):
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for bn in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement('w:' + bn)
            b.set(qn('w:val'), 'none')
            tblBorders.append(b)
        tblPr.append(tblBorders)

    def _p_spacing(p, before=0, after=0):
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)

    def add_secao(titulo):
        p = doc.add_paragraph()
        _p_spacing(p, before=8, after=4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(titulo)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0, 0, 0)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        for side in ('top', 'left', 'bottom', 'right'):
            bdr = OxmlElement('w:' + side)
            bdr.set(qn('w:val'), 'single')
            bdr.set(qn('w:sz'), '4')
            bdr.set(qn('w:space'), '4')
            bdr.set(qn('w:color'), '888888')
            pBdr.append(bdr)
        pPr.append(pBdr)

    def add_dois_col(esq_bold, esq_normal, dir_bold, dir_normal):
        t = doc.add_table(rows=1, cols=2)
        _no_borders(t)
        t.columns[0].width = Mm(110)
        t.columns[1].width = Mm(60)

        lc = t.cell(0, 0)
        lp = lc.paragraphs[0]
        _p_spacing(lp, after=0)
        if esq_bold:
            r = lp.add_run(esq_bold)
            r.bold = True
            r.font.size = Pt(9.5)
        if esq_normal:
            lp2 = lc.add_paragraph()
            _p_spacing(lp2, after=0)
            r2 = lp2.add_run(esq_normal)
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = RGBColor(51, 51, 51)

        rc = t.cell(0, 1)
        rp = rc.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _p_spacing(rp, after=0)
        if dir_bold:
            r3 = rp.add_run(dir_bold)
            r3.bold = True
            r3.font.size = Pt(9)
        if dir_normal:
            rp2 = rc.add_paragraph()
            rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _p_spacing(rp2, after=0)
            r4 = rp2.add_run(dir_normal)
            r4.font.size = Pt(9)
            r4.font.color.rgb = RGBColor(85, 85, 85)

        sp = doc.add_paragraph()
        _p_spacing(sp, after=4)

    def add_bullet(text):
        p = doc.add_paragraph()
        _p_spacing(p, after=2)
        p.paragraph_format.left_indent = Mm(4)
        run = p.add_run('• ' + text)
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(51, 51, 51)

    # ── NOME ──
    p_nome = doc.add_paragraph()
    _p_spacing(p_nome, after=4)
    r_nome = p_nome.add_run(_s(dados.get('nome', 'NOME COMPLETO')).upper())
    r_nome.bold = True
    r_nome.font.size = Pt(17)
    r_nome.font.color.rgb = RGBColor(0, 0, 0)

    # ── CONTATO ──
    linkedin = _s(dados.get('linkedin'))
    github = _s(dados.get('github'))
    email = _s(dados.get('email'))
    telefone = _s(dados.get('telefone'))

    for label, val in (('LinkedIn', linkedin), ('GitHub', github),
                       ('Email', email), ('Celular', telefone)):
        if val:
            p = doc.add_paragraph()
            _p_spacing(p, after=1)
            rb = p.add_run(label + ': ')
            rb.bold = True
            rb.font.size = Pt(9)
            rv = p.add_run(val)
            rv.font.size = Pt(9)
            rv.font.color.rgb = RGBColor(51, 51, 51)

    # divisória
    p_div = doc.add_paragraph()
    _p_spacing(p_div, before=4, after=4)
    pPr = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b_bot = OxmlElement('w:bottom')
    b_bot.set(qn('w:val'), 'single')
    b_bot.set(qn('w:sz'), '8')
    b_bot.set(qn('w:space'), '1')
    b_bot.set(qn('w:color'), '555555')
    pBdr.append(b_bot)
    pPr.append(pBdr)

    # ── FORMAÇÃO ──
    formacao = [f for f in (dados.get('formacao') or [])
                if _s(f.get('instituicao')) or _s(f.get('descricao'))]
    if formacao:
        add_secao('FORMAÇÃO')
        for f in formacao:
            add_dois_col(
                _s(f.get('instituicao')), _s(f.get('descricao')),
                _s(f.get('periodo')), _s(f.get('localidade'))
            )

    # ── COMPETÊNCIAS ──
    comps = [c for c in (dados.get('competencias') or [])
             if _s(c.get('label')) or _s(c.get('valor'))]
    if comps:
        add_secao('RESUMO DE COMPETÊNCIAS')
        for c in comps:
            label = _s(c.get('label'))
            valor = _s(c.get('valor'))
            add_bullet((label + ': ' + valor) if label and valor else (label or valor))

    # ── EXPERIÊNCIA ──
    exps = [e for e in (dados.get('experiencias') or [])
            if _s(e.get('cargo_contrato')) or _s(e.get('empresa'))]
    if exps:
        add_secao('EXPERIÊNCIA PROFISSIONAL')
        for e in exps:
            cargo = _s(e.get('cargo_contrato'))
            empresa = _s(e.get('empresa'))
            periodo = _s(e.get('periodo'))
            resps = _s(e.get('responsabilidades'))
            title_str = cargo + (' | ' + empresa if empresa else '') if cargo else empresa
            add_dois_col(title_str, None, periodo, None)
            if resps:
                bullets = [b.strip().lstrip('•-* ')
                           for b in resps.split('\n') if b.strip()]
                for b in bullets:
                    add_bullet(b)
            sp = doc.add_paragraph()
            _p_spacing(sp, after=4)

    # ── CERTIFICADOS ──
    certs = [c for c in (dados.get('certificados') or []) if _s(c.get('nome'))]
    if certs:
        add_secao('CERTIFICADOS')
        for c in certs:
            add_dois_col('• ' + _s(c.get('nome')), None, _s(c.get('data')), None)

    # ── OUTROS ──
    import base64 as _b64
    from io import BytesIO as _BytesIO

    outros = [o for o in (dados.get('outros') or [])
              if _s(o.get('texto')) or (o.get('imagem') or '').startswith('data:image/')]
    if outros:
        add_secao('OUTROS')
        for o in outros:
            texto = _s(o.get('texto'))
            imagem_b64 = o.get('imagem') or ''
            has_img = imagem_b64.startswith('data:image/')

            if has_img:
                try:
                    raw = imagem_b64.split(',', 1)[1]
                    img_buf = _BytesIO(_b64.b64decode(raw))
                    t = doc.add_table(rows=1, cols=2)
                    _no_borders(t)
                    lc = t.cell(0, 0)
                    if texto:
                        lp = lc.paragraphs[0]
                        _p_spacing(lp, after=2)
                        r = lp.add_run(texto)
                        r.font.size = Pt(9.5)
                        r.font.color.rgb = RGBColor(51, 51, 51)
                    rc = t.cell(0, 1)
                    rp = rc.paragraphs[0]
                    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _p_spacing(rp, after=0)
                    run = rp.add_run()
                    run.add_picture(img_buf, width=Mm(35))
                    sp = doc.add_paragraph()
                    _p_spacing(sp, after=4)
                    continue
                except Exception:
                    pass

            if texto:
                p = doc.add_paragraph()
                _p_spacing(p, after=4)
                r = p.add_run(texto)
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(51, 51, 51)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
